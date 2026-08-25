from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ALLMA_Safety_AI_Project_Pitch.docx"
GOLD = RGBColor(252, 220, 4)
RED = RGBColor(217, 0, 18)
DARK = RGBColor(20, 22, 24)
MUTED = RGBColor(85, 91, 94)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, title, body, accent="FCDC04"):
    shade(cell, "141618")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(accent)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(220, 222, 220)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(15 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    for name in ["Title", "Heading 1", "Heading 2"]:
        doc.styles[name].font.name = "Aptos Display"
        doc.styles[name].font.color.rgb = DARK
    doc.styles["Heading 1"].font.size = Pt(22)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 2"].font.color.rgb = RED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    r = p.add_run("ALLMA SAFETY AI")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RED
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Safety intelligence\nfor real life.")
    r.bold = True; r.font.size = Pt(37); r.font.color.rgb = DARK
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("A conversational, community-powered platform that helps people report, respond, remember, and recover with more clarity.")
    r.font.size = Pt(15); r.font.color.rgb = MUTED
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    set_cell_text(table.cell(0, 0), "ACT NOW", "One calm interface for SOS, incident reporting, nearby help, and trusted response.", "D90012")
    set_cell_text(table.cell(0, 1), "STAY CONNECTED", "Real-time Safety Network coordination with existing in-app voice infrastructure.", "61D39B")
    set_cell_text(table.cell(0, 2), "KEEP LIFE MOVING", "Private Health Reminders for appointments, medication routines, and follow-ups.", "FCDC04")
    p = doc.add_paragraph("Built mobile-first for Uganda. Designed to reduce cognitive load when the moment matters.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True; p.runs[0].font.color.rgb = MUTED
    page_break(doc)

    heading(doc, "The opportunity")
    doc.add_paragraph("When something goes wrong, the next step is rarely obvious.")
    doc.add_paragraph("People need more than a list of emergency numbers. They need a steady layer that turns uncertainty into an actionable path, without pretending to replace human judgment or official services.")
    table = doc.add_table(rows=1, cols=3)
    set_cell_text(table.cell(0, 0), "FRAGMENTED MOMENTS", "Reports, calls, location, contacts, and follow-ups often live in separate tools.", "D90012")
    set_cell_text(table.cell(0, 1), "HIGH COGNITIVE LOAD", "Stress makes complex forms, unclear statuses, and repeated decisions harder to use.", "FCDC04")
    set_cell_text(table.cell(0, 2), "LOW TRUST", "People need to know what is happening, who sees their data, and what happens next.", "7BA7FF")
    heading(doc, "The product insight", 2)
    doc.add_paragraph("ALLMA is a coordination layer: conversational when users need guidance, structured when systems need reliable records, and transparent about what is real.")
    page_break(doc)

    heading(doc, "One assistant. Many moments of safety.")
    doc.add_paragraph("The experience begins with conversation, then moves naturally into the right workflow: a report, a call, a map, a reminder, or a verified community signal.")
    pillars = [("CONVERSATIONAL AI", "Natural-language guidance for crime, missing persons, Lost & Found, nearby facilities, and safety questions."), ("EMERGENCY SOS", "Immediate emergency creation, location state, Safety Network escalation, real response status, and in-app voice."), ("COMMUNITY RESPONSE", "Opt-in responder workflows with authorization, availability, approximate distance, and auditable status changes."), ("HEALTH REMINDERS", "Private appointment, medication, routine, and follow-up reminders with recurrence and explicit notification controls."), ("PUBLIC LOST & FOUND", "Search property handed in to police, submit safe claims, and post lost items for officer matching."), ("POLICE WORKSPACE", "Verified officers review incidents, missing persons, responder activity, claims, and institutional workflows.")]
    table = doc.add_table(rows=2, cols=3)
    for i, (title, body) in enumerate(pillars):
        set_cell_text(table.cell(i // 3, i % 3), title, body, "D90012" if i % 3 == 0 else "FCDC04" if i % 3 == 1 else "61D39B")
    page_break(doc)

    heading(doc, "SOS response: a control screen, not a dashboard")
    doc.add_paragraph("Manual SOS creates the real emergency immediately. The interface answers five questions at a glance: what is happening, who is being contacted, whether location is shared, whether voice is connected, and how to stop or continue.")
    for item in ["Real emergency state is created before location or AI analysis completes.", "Responder states come from actual call and assignment records.", "Location permission never blocks SOS activation.", "A responder can confirm a welfare check; in-app reminders stop only after durable confirmation.", "Emergency calls and ordinary reminder calls remain separate systems."]:
        bullet(doc, item)
    heading(doc, "Response flow", 2)
    flow = doc.add_table(rows=1, cols=4)
    for cell, title, body in zip(flow.rows[0].cells, ["SOS ACTIVE", "LOCATION", "RESPONDERS", "VOICE"], ["Emergency created immediately", "Shared, approximate, or unavailable", "Authorized contacts, real statuses", "ZEGOCLOUD state, never simulated"]):
        set_cell_text(cell, title, body, "D90012" if title == "SOS ACTIVE" else "FCDC04")
    page_break(doc)

    heading(doc, "Health Reminders")
    doc.add_paragraph("A private rhythm for the care people already manage. Health Reminders is optional, calm, and supportive. It does not diagnose, provide medical advice, or make health disclosure part of account creation.")
    for item in ["Create: doctor visits, hospital appointments, follow-ups, medication schedules, routine checks, or other reminders.", "Control: optional health context, recurring schedules, app notifications, Do Not Disturb, and explicit phone-call opt-in.", "Deliver: due reminders are claimed atomically and delivered by a backend scheduler, even when the app is closed.", "Review: complete, reschedule, remove, and revisit reminder history. Missed appointments get a gentle follow-up."]:
        bullet(doc, item,)
    heading(doc, "Privacy by design", 2)
    doc.add_paragraph("Health data is stored in its own private model, protected by row-level ownership policies, and excluded from SOS, Safety Network, community responder, and advertising paths.")
    page_break(doc)

    heading(doc, "Public Lost & Found")
    doc.add_paragraph("A safer bridge between people and police property desks. Anyone can search found property without signing in, submit a claim for officer review, or post something lost so it can be matched.")
    for item in ["Search: free-text search, district filter, and category chips for phones, bags, documents, wallets, keys, and other property.", "Claim safely: masked identifiers, coarse areas, and proof-of-ownership text keep claims meaningful and private.", "Police review: verified officers see contact details, proof, pending claims, and public lost reports in the command workspace.", "No demo claims are approved. Approval is an officer action, recorded with an audit event, and releases the item only after review."]:
        bullet(doc, item)
    page_break(doc)

    heading(doc, "Trust architecture")
    doc.add_paragraph("ALLMA uses AI to assist understanding, not to invent certainty. It uses automation to keep a process moving, not to hide decisions.")
    for item in ["Supabase-backed records with row-level security and explicit ownership boundaries.", "Verified-officer policies for police-side review and public-facing safe summaries.", "Authorized Safety Network matching, consent-aware location sharing, and real call state.", "ZEGOCLOUD remains the in-app emergency voice layer; ordinary reminder calls are separate.", "Backend scheduling uses atomic due-claiming to prevent duplicate reminder delivery.", "Honest states for unavailable location, weak connection, no answer, and failed delivery."]:
        bullet(doc, item)
    page_break(doc)

    heading(doc, "Current build and roadmap")
    doc.add_paragraph("The current project is implemented as a mobile-first TanStack Start application with Supabase persistence and focused workflow surfaces.")
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.cell(0, 0), "IMPLEMENTED NOW", "Conversational safety assistant\nSOS activation and live response UI\nSafety Network and responder workflows\nZEGOCLOUD in-app voice\nPrivate Health Reminders\nPublic Lost & Found\nResponsive dark/light UI", "61D39B")
    set_cell_text(table.cell(0, 1), "NEXT HARDENING", "Production notification configuration\nOperational monitoring\nExpanded officer matching\nNative background delivery\nPilot measurement with communities and partners", "FCDC04")
    heading(doc, "Roadmap", 2)
    for item in ["Pilot: validate answer rates, completion, false alarms, and user comprehension.", "Reliability: harden push, scheduler, reconnection, background behavior, and operational alerting.", "Network: expand verified partners, facilities, police property workflows, and responder coverage.", "Intelligence: improve routing and guidance with privacy-preserving signals, without profiling or health inference."]:
        bullet(doc, item)
    page_break(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    r = p.add_run("Help us make safety more understandable.")
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = DARK
    p = doc.add_paragraph("ALLMA is ready for the next step: responsible pilots, trusted institutional partners, and the operational support required to turn a strong product foundation into dependable public infrastructure.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14); p.runs[0].font.color.rgb = MUTED
    table = doc.add_table(rows=1, cols=3)
    set_cell_text(table.cell(0, 0), "PARTNER", "Community organisations, police property desks, health facilities, and responders.", "61D39B")
    set_cell_text(table.cell(0, 1), "PILOT", "Test the workflows with real people, real constraints, and measurable outcomes.", "FCDC04")
    set_cell_text(table.cell(0, 2), "BUILD", "Support reliability, safety operations, and platform expansion.", "D90012")
    p = doc.add_paragraph("No traction, coverage, or response-time figures are claimed until they are measured in the field.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True; p.runs[0].font.color.rgb = MUTED
    doc.save(OUT)
    print(OUT)

if __name__ == "__main__":
    build()
